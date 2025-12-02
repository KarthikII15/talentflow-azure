# # backend/main.py
# from dotenv import load_dotenv
# import os
# from fastapi import FastAPI, Query, HTTPException
# from pydantic import BaseModel
# from typing import Literal
# import uuid
# from io import BytesIO

# # ⭐ ADD CORS ⭐
# from fastapi.middleware.cors import CORSMiddleware

# load_dotenv()

# # ---------- INTERNAL IMPORTS ----------
# from backend.db.cloud_provider import AWSProvider, GCPProvider, AzureProvider, CloudProvider
# from backend.agents.recruiter import evaluate_candidate
# from backend.agents.cloud_setup import CloudSetupAgent
# from backend.agents.librarian import LibrarianAgent

# from backend.db.firestore_client import (
#     save_jd, save_candidate, list_jds, list_candidates,
#     save_mapping, list_mappings_for_jd,
#     get_candidate, get_jd, save_shortlist_result
# )

# # ---------- EXTERNAL IMPORTS ----------
# from google.cloud import storage
# from pdfminer.high_level import extract_text as pdf_extract_text
# from docx import Document


# # ---------------------------------------------------------
# # FASTAPI APP INITIALIZATION
# # ---------------------------------------------------------
# app = FastAPI(title="CARPAS – Cloud-Agnostic Recruitment Pipeline Automation System")

# # ⭐⭐ IMPORTANT: ENABLE CORS ⭐⭐
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=[
#         "http://localhost:5173",
#         "http://127.0.0.1:5173",
#     ],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )


# # ---------------------------------------------------------
# # MODELS
# # ---------------------------------------------------------
# CloudName = Literal["aws", "gcp", "azure"]


# class CandidateEvalRequest(BaseModel):
#     job_id: str
#     candidate_id: str
#     required_experience: float
#     candidate_experience: float
#     career_gaps_detected: bool
#     job_shifts_last_year: int
#     skills_match_score: int


# class MappingRequest(BaseModel):
#     jd_id: str
#     candidate_id: str
#     status: Literal["PENDING", "SHORTLISTED", "REJECTED"] = "PENDING"


# class ShortlistRequest(BaseModel):
#     candidate_id: str
#     skills_match_score: int = 70
    
# class UploadURLRequest(BaseModel):
#     filename: str


# # ---------------------------------------------------------
# # CLOUD PROVIDER FACTORY
# # ---------------------------------------------------------
# def get_provider(provider: CloudName) -> CloudProvider:
#     if provider == "gcp":
#         return GCPProvider()
#     if provider == "aws":
#         return AWSProvider()
#     if provider == "azure":
#         return AzureProvider()
#     raise ValueError("Invalid provider")


# # ---------------------------------------------------------
# # ROUTES
# # ---------------------------------------------------------

# @app.get("/providers")
# def list_providers():
#     return {"providers": ["aws", "gcp", "azure"]}


# @app.get("/cloud/setup")
# def setup_cloud(provider: CloudName = Query(...)):
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     provider_obj = get_provider("gcp")
#     agent = CloudSetupAgent(provider_obj)
#     return agent.setup_storage()

# @app.post("/cloud/generate-upload-url")
# def generate_upload_url(payload: UploadURLRequest):
#     bucket_name = "talentflow-gcp"
#     path = f"incoming/{payload.filename}"

#     client = storage.Client()
#     bucket = client.bucket(bucket_name)
#     blob = bucket.blob(path)

#     url = blob.generate_signed_url(
#         version="v4",
#         expiration=3600,  # valid 1 hour
#         method="PUT",
#         content_type="application/octet-stream",
#     )

#     return {"upload_url": url}


# @app.post("/candidate/evaluate")
# def evaluate_candidate_api(payload: CandidateEvalRequest):
#     return evaluate_candidate(
#         job_id=payload.job_id,
#         candidate_id=payload.candidate_id,
#         required_experience=payload.required_experience,
#         candidate_experience=payload.candidate_experience,
#         career_gaps_detected=payload.career_gaps_detected,
#         job_shifts_last_year=payload.job_shifts_last_year,
#         skills_match_score=payload.skills_match_score,
#     )


# # ---------------------------------------------------------
# # TEXT EXTRACTION MODULE
# # ---------------------------------------------------------
# def extract_text_from_blob(bucket: storage.Bucket, blob_name: str) -> str:
#     ext = blob_name.lower().split(".")[-1]
#     blob = bucket.blob(blob_name)

#     if ext in ("txt", "md", "log"):
#         return blob.download_as_text(errors="ignore")

#     if ext == "pdf":
#         try:
#             return pdf_extract_text(BytesIO(blob.download_as_bytes()))
#         except Exception as e:
#             print("PDF extract error:", e)
#             return ""

#     if ext == "docx":
#         try:
#             doc = Document(BytesIO(blob.download_as_bytes()))
#             return "\n".join(p.text for p in doc.paragraphs)
#         except Exception as e:
#             print("DOCX extract error:", e)
#             return ""

#     return ""


# # ---------------------------------------------------------
# # FILE PROCESSING
# # ---------------------------------------------------------
# @app.post("/cloud/process-files")
# def process_files(provider: CloudName = Query(...)):
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     bucket_name = "talentflow-gcp"
#     gcp_client = storage.Client()
#     bucket = gcp_client.bucket(bucket_name)
#     provider_obj = get_provider("gcp")
#     librarian = LibrarianAgent()

#     prefix = "incoming/"
#     file_names = provider_obj.list_files(prefix)

#     if not file_names:
#         return {
#             "message": "No files found",
#             "total_files": 0,
#             "resumes": 0,
#             "jds": 0,
#             "others": 0,
#             "unknown": 0,
#             "details": [],
#         }

#     details = []
#     resumes_count = jds_count = others_count = unknown_count = 0

#     for name in file_names:
#         text = extract_text_from_blob(bucket, name)
#         filename_only = name.split("/")[-1]

#         if not text.strip():
#             unknown_count += 1
#             provider_obj.move_file(name, f"Others/{filename_only}")
#             details.append({
#                 "file": name,
#                 "type": "UNKNOWN",
#                 "moved_to": f"Others/{filename_only}",
#                 "status": "MOVED"
#             })
#             continue

#         doc_type = librarian.classify_document(text)

#         if doc_type == "RESUME":
#             candidate_id = "CAND-" + str(uuid.uuid4())[:8]
#             save_candidate(candidate_id, filename_only, text)
#             dest = f"Resumes/{filename_only}"
#             resumes_count += 1

#         elif doc_type == "JD":
#             jd_id = "JD-" + str(uuid.uuid4())[:8]
#             save_jd(jd_id, filename_only, text)
#             dest = f"JDs/{filename_only}"
#             jds_count += 1

#         else:
#             dest = f"Others/{filename_only}"
#             others_count += 1

#         provider_obj.move_file(name, dest)

#         details.append({
#             "file": name,
#             "type": doc_type,
#             "moved_to": dest,
#             "status": "MOVED"
#         })

#     return {
#         "message": "Processing completed.",
#         "bucket": bucket_name,
#         "total_files": len(file_names),
#         "resumes": resumes_count,
#         "jds": jds_count,
#         "others": others_count,
#         "unknown": unknown_count,
#         "details": details,
#     }


# # ---------------------------------------------------------
# # LIST ENDPOINTS (SAFE VERSION)
# # ---------------------------------------------------------
# @app.get("/jds")
# def get_jds():
#     try:
#         return {"jds": list_jds()}
#     except Exception as e:
#         print("ERROR /jds:", e)
#         return {"jds": []}


# @app.get("/candidates")
# def get_candidates():
#     try:
#         return {"candidates": list_candidates()}
#     except Exception as e:
#         print("ERROR /candidates:", e)
#         return {"candidates": []}


# # ---------------------------------------------------------
# # JD <-> CANDIDATE MAPPING
# # ---------------------------------------------------------
# @app.post("/map-candidate-to-jd")
# def map_candidate_to_jd(payload: MappingRequest):
#     if not get_jd(payload.jd_id):
#         raise HTTPException(404, "JD not found")
#     if not get_candidate(payload.candidate_id):
#         raise HTTPException(404, "Candidate not found")

#     save_mapping(payload.jd_id, payload.candidate_id, payload.status)
#     return {"message": "Mapping saved"}


# @app.get("/jd/{jd_id}/candidates")
# def get_candidates_for_jd(jd_id: str):
#     mappings = list_mappings_for_jd(jd_id)
#     candidate_ids = [m["candidate_id"] for m in mappings]
#     all_candidates = {c["candidate_id"]: c for c in list_candidates()}

#     linked = [all_candidates[cid] for cid in candidate_ids if cid in all_candidates]

#     return {"jd_id": jd_id, "candidates": linked, "mappings": mappings}


# # ---------------------------------------------------------
# # SHORTLISTING ENGINE
# # ---------------------------------------------------------
# @app.post("/jd/{jd_id}/shortlist")
# def shortlist_candidate(jd_id: str, payload: ShortlistRequest):
#     jd = get_jd(jd_id)
#     cand = get_candidate(payload.candidate_id)

#     if not jd or not cand:
#         raise HTTPException(404, "JD or Candidate not found")

#     eval_result = evaluate_candidate(
#         job_id=jd_id,
#         candidate_id=payload.candidate_id,
#         required_experience=3.0,
#         candidate_experience=3.0,
#         career_gaps_detected=False,
#         job_shifts_last_year=0,
#         skills_match_score=payload.skills_match_score,
#     )

#     save_shortlist_result(jd_id, payload.candidate_id, eval_result)
#     return eval_result


# # backend/main.py
# from dotenv import load_dotenv
# import os
# from fastapi import FastAPI, Query, HTTPException
# from pydantic import BaseModel
# from typing import Literal
# import uuid
# from io import BytesIO

# from fastapi.middleware.cors import CORSMiddleware

# load_dotenv()

# # ---------- INTERNAL IMPORTS ----------
# from backend.db.cloud_provider import AWSProvider, GCPProvider, AzureProvider, CloudProvider
# from backend.agents.recruiter import evaluate_candidate
# from backend.agents.cloud_setup import CloudSetupAgent
# from backend.agents.librarian import LibrarianAgent

# from backend.db.firestore_client import (
#     save_jd, save_candidate, list_jds, list_candidates,
#     save_mapping, list_mappings_for_jd,
#     get_candidate, get_jd, save_shortlist_result
# )

# # ---------- EXTERNAL IMPORTS ----------
# from google.cloud import storage
# from pdfminer.high_level import extract_text as pdf_extract_text
# from docx import Document

# # ---------------------------------------------------------
# # FASTAPI APP INITIALIZATION
# # ---------------------------------------------------------
# app = FastAPI(title="CARPAS – Cloud-Agnostic Recruitment Pipeline Automation System")

# # Enable CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # ---------------------------------------------------------
# # MODELS
# # ---------------------------------------------------------
# CloudName = Literal["aws", "gcp", "azure"]


# class CandidateEvalRequest(BaseModel):
#     job_id: str
#     candidate_id: str
#     required_experience: float
#     candidate_experience: float
#     career_gaps_detected: bool
#     job_shifts_last_year: int
#     skills_match_score: int


# class MappingRequest(BaseModel):
#     jd_id: str
#     candidate_id: str
#     status: Literal["PENDING", "SHORTLISTED", "REJECTED"] = "PENDING"


# class ShortlistRequest(BaseModel):
#     candidate_id: str
#     skills_match_score: int = 70


# class UploadURLRequest(BaseModel):
#     filename: str


# # ---------------------------------------------------------
# # CLOUD PROVIDER FACTORY
# # ---------------------------------------------------------
# def get_provider(provider: CloudName) -> CloudProvider:
#     if provider == "gcp":
#         return GCPProvider()
#     if provider == "aws":
#         return AWSProvider()
#     if provider == "azure":
#         return AzureProvider()
#     raise ValueError("Invalid provider")


# # ---------------------------------------------------------
# # ROUTES
# # ---------------------------------------------------------
# @app.get("/providers")
# def list_providers():
#     return {"providers": ["aws", "gcp", "azure"]}


# @app.get("/cloud/setup")
# def setup_cloud(provider: CloudName = Query(...)):
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     provider_obj = get_provider("gcp")
#     agent = CloudSetupAgent(provider_obj)
#     return agent.setup_storage()


# # ---------------------------------------------------------
# # PER-FILE SIGNED UPLOAD URL
# # ---------------------------------------------------------
# @app.post("/cloud/generate-upload-url")
# def generate_upload_url(payload: UploadURLRequest):
#     bucket_name = "talentflow-gcp"

#     provider = GCPProvider()

#     # 👇 The object path MUST be a file path
#     object_path = f"incoming/{payload.filename}"

#     # 👇 Generate signed URL for THIS EXACT file
#     bucket = provider._get_bucket()
#     blob = bucket.blob(object_path)

#     signed_url = blob.generate_signed_url(
#         version="v4",
#         expiration=3600,
#         method="PUT",
#         content_type="application/octet-stream",
#     )
#     print("SIGNED URL:", signed_url)
#     return {"upload_url": signed_url}



# # ---------------------------------------------------------
# # CANDIDATE EVALUATION
# # ---------------------------------------------------------
# @app.post("/candidate/evaluate")
# def evaluate_candidate_api(payload: CandidateEvalRequest):
#     return evaluate_candidate(
#         job_id=payload.job_id,
#         candidate_id=payload.candidate_id,
#         required_experience=payload.required_experience,
#         candidate_experience=payload.candidate_experience,
#         career_gaps_detected=payload.career_gaps_detected,
#         job_shifts_last_year=payload.job_shifts_last_year,
#         skills_match_score=payload.skills_match_score,
#     )


# # ---------------------------------------------------------
# # TEXT EXTRACTION
# # ---------------------------------------------------------
# def extract_text_from_blob(bucket: storage.Bucket, blob_name: str) -> str:
#     ext = blob_name.lower().split(".")[-1]
#     blob = bucket.blob(blob_name)

#     if ext in ("txt", "md", "log"):
#         return blob.download_as_text(errors="ignore")

#     if ext == "pdf":
#         try:
#             return pdf_extract_text(BytesIO(blob.download_as_bytes()))
#         except:
#             return ""

#     if ext == "docx":
#         try:
#             doc = Document(BytesIO(blob.download_as_bytes()))
#             return "\n".join(p.text for p in doc.paragraphs)
#         except:
#             return ""

#     return ""


# # ---------------------------------------------------------
# # FILE PROCESSING PIPELINE
# # ---------------------------------------------------------
# @app.post("/cloud/process-files")
# def process_files(provider: CloudName = Query(...)):
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     bucket_name = "talentflow-gcp"
#     gcp_client = storage.Client()
#     bucket = gcp_client.bucket(bucket_name)
#     provider_obj = get_provider("gcp")
#     librarian = LibrarianAgent()

#     prefix = "incoming/"
#     file_names = provider_obj.list_files(prefix)

#     if not file_names:
#         return {
#             "message": "No files found",
#             "total_files": 0,
#             "resumes": 0,
#             "jds": 0,
#             "others": 0,
#             "unknown": 0,
#             "details": [],
#         }

#     details = []
#     resumes_count = jds_count = others_count = unknown_count = 0

#     for name in file_names:
#         text = extract_text_from_blob(bucket, name)
#         filename_only = name.split("/")[-1]

#         if not text.strip():
#             unknown_count += 1
#             provider_obj.move_file(name, f"Others/{filename_only}")
#             details.append({
#                 "file": name,
#                 "type": "UNKNOWN",
#                 "moved_to": f"Others/{filename_only}",
#                 "status": "MOVED"
#             })
#             continue

#         doc_type = librarian.classify_document(text)

#         if doc_type == "RESUME":
#             candidate_id = "CAND-" + str(uuid.uuid4())[:8]
#             save_candidate(candidate_id, filename_only, text)
#             dest = f"Resumes/{filename_only}"
#             resumes_count += 1

#         elif doc_type == "JD":
#             jd_id = "JD-" + str(uuid.uuid4())[:8]
#             save_jd(jd_id, filename_only, text)
#             dest = f"JDs/{filename_only}"
#             jds_count += 1

#         else:
#             dest = f"Others/{filename_only}"
#             others_count += 1

#         provider_obj.move_file(name, dest)

#         details.append({
#             "file": name,
#             "type": doc_type,
#             "moved_to": dest,
#             "status": "MOVED"
#         })

#     return {
#         "message": "Processing completed.",
#         "bucket": bucket_name,
#         "total_files": len(file_names),
#         "resumes": resumes_count,
#         "jds": jds_count,
#         "others": others_count,
#         "unknown": unknown_count,
#         "details": details,
#     }


# # ---------------------------------------------------------
# # SAFE LIST ROUTES
# # ---------------------------------------------------------
# @app.get("/jds")
# def get_jds():
#     try:
#         return {"jds": list_jds()}
#     except:
#         return {"jds": []}


# @app.get("/candidates")
# def get_candidates():
#     try:
#         return {"candidates": list_candidates()}
#     except:
#         return {"candidates": []}


# # ---------------------------------------------------------
# # JD <-> CANDIDATE MAPPING
# # ---------------------------------------------------------
# @app.post("/map-candidate-to-jd")
# def map_candidate_to_jd(payload: MappingRequest):
#     if not get_jd(payload.jd_id):
#         raise HTTPException(404, "JD not found")
#     if not get_candidate(payload.candidate_id):
#         raise HTTPException(404, "Candidate not found")

#     save_mapping(payload.jd_id, payload.candidate_id, payload.status)
#     return {"message": "Mapping saved"}


# @app.get("/jd/{jd_id}/candidates")
# def get_candidates_for_jd(jd_id: str):
#     mappings = list_mappings_for_jd(jd_id)
#     candidate_ids = [m["candidate_id"] for m in mappings]
#     all_candidates = {c["candidate_id"]: c for c in list_candidates()}
#     linked = [all_candidates[cid] for cid in candidate_ids if cid in all_candidates]

#     return {"jd_id": jd_id, "candidates": linked, "mappings": mappings}


# # ---------------------------------------------------------
# # SHORTLISTING ENGINE
# # ---------------------------------------------------------
# @app.post("/jd/{jd_id}/shortlist")
# def shortlist_candidate(jd_id: str, payload: ShortlistRequest):
#     jd = get_jd(jd_id)
#     cand = get_candidate(payload.candidate_id)

#     if not jd or not cand:
#         raise HTTPException(404, "JD or Candidate not found")

#     eval_result = evaluate_candidate(
#         job_id=jd_id,
#         candidate_id=payload.candidate_id,
#         required_experience=3.0,
#         candidate_experience=3.0,
#         career_gaps_detected=False,
#         job_shifts_last_year=0,
#         skills_match_score=payload.skills_match_score,
#     )

#     save_shortlist_result(jd_id, payload.candidate_id, eval_result)
#     return eval_result


# from dotenv import load_dotenv
# import os
# import uuid
# import traceback
# from io import BytesIO
# from typing import Literal
# from datetime import timedelta

# from fastapi import FastAPI, Query, HTTPException, Header
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel

# # ---------- EXTERNAL IMPORTS ----------
# from google.cloud import storage
# from pdfminer.high_level import extract_text as pdf_extract_text
# from docx import Document

# # ---------- INTERNAL IMPORTS ----------
# from backend.db.cloud_provider import AWSProvider, GCPProvider, AzureProvider, CloudProvider
# from backend.agents.recruiter import evaluate_candidate
# from backend.agents.cloud_setup import CloudSetupAgent
# from backend.agents.librarian import LibrarianAgent

# from backend.db.firestore_client import (
#     save_jd, save_candidate, list_jds, list_candidates,
#     save_mapping, list_mappings_for_jd,
#     get_candidate, get_jd, save_shortlist_result
# )

# load_dotenv()

# # ---------------------------------------------------------
# # FASTAPI APP INITIALIZATION
# # ---------------------------------------------------------
# app = FastAPI(title="CARPAS – Cloud-Agnostic Recruitment Pipeline Automation System")

# # Enable CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # ---------------------------------------------------------
# # MODELS
# # ---------------------------------------------------------
# CloudName = Literal["aws", "gcp", "azure"]

# class CandidateEvalRequest(BaseModel):
#     job_id: str
#     candidate_id: str
#     required_experience: float
#     candidate_experience: float
#     career_gaps_detected: bool
#     job_shifts_last_year: int
#     skills_match_score: int

# class MappingRequest(BaseModel):
#     jd_id: str
#     candidate_id: str
#     status: Literal["PENDING", "SHORTLISTED", "REJECTED"] = "PENDING"

# class ShortlistRequest(BaseModel):
#     candidate_id: str
#     skills_match_score: int = 70

# class UploadURLRequest(BaseModel):
#     filename: str

# # ---------------------------------------------------------
# # HELPER FUNCTIONS
# # ---------------------------------------------------------
# def get_provider(provider: CloudName) -> CloudProvider:
#     if provider == "gcp":
#         return GCPProvider()
#     if provider == "aws":
#         return AWSProvider()
#     if provider == "azure":
#         return AzureProvider()
#     raise ValueError("Invalid provider")

# def extract_text_from_blob(bucket: storage.Bucket, blob_name: str) -> str:
#     """
#     Downloads file bytes from GCS and extracts text based on extension.
#     """
#     try:
#         blob = bucket.blob(blob_name)
#         # Read file extension
#         ext = blob_name.lower().split(".")[-1]
        
#         # Download bytes into memory
#         data = blob.download_as_bytes()
        
#         if ext == "pdf":
#             return pdf_extract_text(BytesIO(data))
#         elif ext == "docx":
#             doc = Document(BytesIO(data))
#             return "\n".join([p.text for p in doc.paragraphs])
#         elif ext in ["txt", "md", "log"]:
#             return data.decode("utf-8", errors="ignore")
#         else:
#             return ""
#     except Exception as e:
#         print(f"Error extracting text from {blob_name}: {e}")
#         return ""

# # ---------------------------------------------------------
# # INFRASTRUCTURE ROUTES
# # ---------------------------------------------------------
# @app.get("/providers")
# def list_providers():
#     return {"providers": ["aws", "gcp", "azure"]}

# @app.get("/cloud/setup")
# def setup_cloud(provider: CloudName = Query(...)):
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     provider_obj = get_provider("gcp")
#     agent = CloudSetupAgent(provider_obj)
#     return agent.setup_storage()

# # ---------------------------------------------------------
# # INGESTION ROUTES
# # ---------------------------------------------------------
# @app.post("/cloud/generate-upload-url")
# def generate_upload_url(payload: UploadURLRequest):
#     try:
#         provider = GCPProvider() # Defaults to talentflow-gcp bucket
        
#         # We place files in 'incoming/' virtually
#         object_path = f"incoming/{payload.filename}"
        
#         bucket = provider._get_bucket()
#         blob = bucket.blob(object_path)

#         # ⬇️ No content_type enforced here to avoid signature mismatch with frontend
#         signed_url = blob.generate_signed_url(
#             version="v4",
#             expiration=timedelta(hours=1),
#             method="PUT",
#         )

#         print(f"GENERATED URL FOR: {object_path}")
#         return {
#             "upload_url": signed_url,
#             "object_path": object_path,
#         }
#     except Exception as e:
#         print("ERROR generating upload URL:", e)
#         raise HTTPException(status_code=500, detail=f"Failed to generate upload URL: {e}")

# # ---------------------------------------------------------
# # PIPELINE ROUTE (The Logic You Requested)
# # ---------------------------------------------------------
# @app.post("/cloud/process-files")
# def process_files(provider: CloudName = Query(...)):
#     try:
#         # 1. Setup Provider & Bucket
#         provider_obj = get_provider(provider)
#         bucket = provider_obj._get_bucket()

#         # 2. List incoming files
#         # Note: Ensure your list_files implementation filters out the folder placeholder itself "incoming/"
#         incoming_files = provider_obj.list_files("incoming/")

#         results = {
#             "bucket": provider_obj.bucket_name,
#             "total_files": len(incoming_files),
#             "resumes": 0,
#             "jds": 0,
#             "others": 0,
#             "unknown": 0,
#             "skipped": 0,
#             "details": [],
#         }

#         # 3. Initialize AI Agent
#         librarian = LibrarianAgent()

#         for file_path in incoming_files:
#             filename = file_path.split("/")[-1]
#             if not filename: continue # Skip if it's just the folder path

#             # Extract Text
#             text = extract_text_from_blob(bucket, file_path)

#             if not text.strip():
#                 results["skipped"] += 1
#                 # Move empty files to Others so they don't block the queue
#                 provider_obj.move_file(file_path, f"Others/{filename}")
#                 continue

#             # AI CLASSIFIER
#             # The LibrarianAgent returns a string (e.g., "RESUME", "JD")
#             # If your agent returns a dict, adjust below. Assuming it returns string based on your agent code.
#             doc_type = librarian.classify_document(text) 
#             # classification = librarian.classify_document(text)
#             # doc_type = classification.get("document_type", "UNKNOWN")


#             dest = ""
            
#             if doc_type == "RESUME":
#                 dest = f"Resumes/{filename}"
#                 # Generate ID for DB
#                 cand_id = f"CAND-{uuid.uuid4().hex[:8]}"
#                 save_candidate(cand_id, filename, text)
#                 results["resumes"] += 1

#             elif doc_type == "JD":
#                 dest = f"JDs/{filename}"
#                 # Generate ID for DB
#                 jd_id = f"JD-{uuid.uuid4().hex[:8]}"
#                 save_jd(jd_id, filename, text)
#                 results["jds"] += 1

#             else:
#                 dest = f"Others/{filename}"
#                 results["others"] += 1

#             # Move file in Cloud Storage
#             provider_obj.move_file(file_path, dest)

#             results["details"].append({
#                 "file": file_path,
#                 "type": doc_type,
#                 "moved_to": dest,
#                 "status": "MOVED"
#             })

#         return results

#     except Exception as e:
#         print("PROCESS ERROR:")
#         traceback.print_exc() # Print full stack trace to console for debugging
#         raise HTTPException(status_code=500, detail=f"Error processing files: {str(e)}")

# # ---------------------------------------------------------
# # DATA ROUTES
# # ---------------------------------------------------------
# @app.get("/jds")
# def get_jds_route():
#     try:
#         return {"jds": list_jds()}
#     except:
#         return {"jds": []}

# @app.get("/candidates")
# def get_candidates_route():
#     try:
#         return {"candidates": list_candidates()}
#     except:
#         return {"candidates": []}

# # ---------------------------------------------------------
# # MAPPING ROUTES
# # ---------------------------------------------------------
# @app.post("/map-candidate-to-jd")
# def map_candidate_to_jd(payload: MappingRequest):
#     if not get_jd(payload.jd_id):
#         raise HTTPException(404, "JD not found")
#     if not get_candidate(payload.candidate_id):
#         raise HTTPException(404, "Candidate not found")

#     save_mapping(payload.jd_id, payload.candidate_id, payload.status)
#     return {"message": "Mapping saved"}

# @app.get("/jd/{jd_id}/candidates")
# def get_candidates_for_jd(jd_id: str):
#     mappings = list_mappings_for_jd(jd_id)
#     # Get all candidates
#     all_candidates = {c["candidate_id"]: c for c in list_candidates()}
    
#     # Filter only those mapped to this JD
#     linked = []
#     for m in mappings:
#         cid = m["candidate_id"]
#         if cid in all_candidates:
#             linked.append(all_candidates[cid])

#     return {"jd_id": jd_id, "candidates": linked, "mappings": mappings}

# # ---------------------------------------------------------
# # EVALUATION ROUTES
# # ---------------------------------------------------------
# @app.post("/candidate/evaluate")
# def evaluate_candidate_api(payload: CandidateEvalRequest):
#     return evaluate_candidate(
#         job_id=payload.job_id,
#         candidate_id=payload.candidate_id,
#         required_experience=payload.required_experience,
#         candidate_experience=payload.candidate_experience,
#         career_gaps_detected=payload.career_gaps_detected,
#         job_shifts_last_year=payload.job_shifts_last_year,
#         skills_match_score=payload.skills_match_score,
#     )

# @app.post("/jd/{jd_id}/shortlist")
# def shortlist_candidate(jd_id: str, payload: ShortlistRequest):
#     jd = get_jd(jd_id)
#     cand = get_candidate(payload.candidate_id)

#     if not jd or not cand:
#         raise HTTPException(404, "JD or Candidate not found")

#     # In a real scenario, you'd extract these numbers from the Resume/JD text analysis
#     # For now, we use defaults or mock logic
#     eval_result = evaluate_candidate(
#         job_id=jd_id,
#         candidate_id=payload.candidate_id,
#         required_experience=3.0,
#         candidate_experience=3.0,
#         career_gaps_detected=False,
#         job_shifts_last_year=0,
#         skills_match_score=payload.skills_match_score,
#     )

#     save_shortlist_result(
    # 
    # jd_id, payload.candidate_id, eval_result)
#     return eval_result

# from dotenv import load_dotenv
# import os
# import uuid
# import traceback
# from io import BytesIO
# from typing import Literal
# from datetime import timedelta

# from fastapi import FastAPI, Query, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel

# # ---------- EXTERNAL IMPORTS ----------
# from google.cloud import storage
# from pdfminer.high_level import extract_text as pdf_extract_text
# from docx import Document

# # ---------- INTERNAL IMPORTS ----------
# from backend.db.cloud_provider import (
#     AWSProvider,
#     GCPProvider,
#     AzureProvider,
#     CloudProvider,
# )
# from backend.agents.recruiter import evaluate_candidate
# from backend.agents.cloud_setup import CloudSetupAgent
# from backend.agents.librarian import LibrarianAgent

# from backend.db.firestore_client import (
#     save_jd,
#     save_candidate,
#     list_jds,
#     list_candidates,
#     save_mapping,
#     list_mappings_for_jd,
#     get_candidate,
#     get_jd,
#     save_shortlist_result,
# )

# load_dotenv()

# # ---------------------------------------------------------
# # FASTAPI APP INITIALIZATION
# # ---------------------------------------------------------
# app = FastAPI(title="CARPAS – Cloud-Agnostic Recruitment Pipeline Automation System")

# # Enable CORS for Vite dev server
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=[
#         "http://localhost:5173",
#         "http://127.0.0.1:5173",
#     ],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # ---------------------------------------------------------
# # MODELS
# # ---------------------------------------------------------
# CloudName = Literal["aws", "gcp", "azure"]


# class CandidateEvalRequest(BaseModel):
#     job_id: str
#     candidate_id: str
#     required_experience: float
#     candidate_experience: float
#     career_gaps_detected: bool
#     job_shifts_last_year: int
#     skills_match_score: int


# class MappingRequest(BaseModel):
#     jd_id: str
#     candidate_id: str
#     status: Literal["PENDING", "SHORTLISTED", "REJECTED"] = "PENDING"


# class ShortlistRequest(BaseModel):
#     candidate_id: str
#     skills_match_score: int = 70


# class UploadURLRequest(BaseModel):
#     filename: str


# # ---------------------------------------------------------
# # HELPERS
# # ---------------------------------------------------------
# def get_provider(provider: CloudName) -> CloudProvider:
#     if provider == "gcp":
#         return GCPProvider()
#     if provider == "aws":
#         return AWSProvider()
#     if provider == "azure":
#         return AzureProvider()
#     raise ValueError("Invalid provider")


# def extract_text_from_blob(bucket: storage.Bucket, blob_name: str) -> str:
#     """
#     Download file from GCS and extract text depending on extension.
#     """
#     try:
#         blob = bucket.blob(blob_name)
#         data = blob.download_as_bytes()
#         ext = blob_name.lower().split(".")[-1]

#         if ext == "pdf":
#             return pdf_extract_text(BytesIO(data))

#         if ext == "docx":
#             doc = Document(BytesIO(data))
#             return "\n".join(p.text for p in doc.paragraphs)

#         if ext in ("txt", "md", "log"):
#             return data.decode("utf-8", errors="ignore")

#         # Unsupported extension → treat as empty
#         return ""
#     except Exception as e:
#         print(f"[TEXT EXTRACTION ERROR] {blob_name} → {e}")
#         return ""


# # ---------------------------------------------------------
# # INFRASTRUCTURE ROUTES
# # ---------------------------------------------------------
# @app.get("/providers")
# def list_providers():
#     return {"providers": ["aws", "gcp", "azure"]}


# @app.get("/cloud/setup")
# def setup_cloud(provider: CloudName = Query(...)):
#     # For now only GCP is implemented
#     if provider != "gcp":
#         raise HTTPException(400, "Only GCP implemented.")

#     provider_obj = get_provider("gcp")
#     agent = CloudSetupAgent(provider_obj)
#     return agent.setup_storage()


# # ---------------------------------------------------------
# # INGESTION ROUTES
# # ---------------------------------------------------------
# @app.post("/cloud/generate-upload-url")
# def generate_upload_url(payload: UploadURLRequest):
#     """
#     Frontend calls this once per file to get a signed URL.
#     Files are uploaded directly from browser -> GCS 'incoming/'.
#     """
#     try:
#         provider = GCPProvider()  # uses talentflow-gcp by default

#         object_path = f"incoming/{payload.filename}"
#         bucket = provider._get_bucket()
#         blob = bucket.blob(object_path)

#         # No content_type -> keep URL simple and avoid header mismatch
#         signed_url = blob.generate_signed_url(
#             version="v4",
#             expiration=timedelta(hours=1),
#             method="PUT",
#         )

#         print(f"[SIGNED URL] {object_path}")
#         return {
#             "upload_url": signed_url,
#             "object_path": object_path,
#         }
#     except Exception as e:
#         print("[ERROR] generating upload URL:", e)
#         raise HTTPException(
#             status_code=500, detail=f"Failed to generate upload URL: {e}"
#         )


# # ---------------------------------------------------------
# # AI CLASSIFICATION PIPELINE
# # ---------------------------------------------------------
# @app.post("/cloud/process-files")
# def process_files(provider: CloudName = Query(...)):
#     """
#     Main pipeline:
#     - Reads all files from 'incoming/' in the selected cloud
#     - Extracts text
#     - Classifies as RESUME / JD / OTHER using LibrarianAgent
#     - Saves structured text into Firestore
#     - Moves files into JDs/, Resumes/, Others/
#     """
#     try:
#         # 1. Only GCP is implemented today
#         if provider != "gcp":
#             raise HTTPException(400, "Only GCP implemented.")

#         provider_obj = get_provider("gcp")
#         bucket = provider_obj._get_bucket()

#         # 2. List all files inside incoming/
#         incoming_files = provider_obj.list_files("incoming/")

#         results = {
#             "message": "Processing completed.",
#             "bucket": provider_obj.bucket_name,
#             "total_files": len(incoming_files),
#             "resumes": 0,
#             "jds": 0,
#             "others": 0,
#             "unknown": 0,
#             "skipped": 0,
#             "details": [],
#         }

#         if not incoming_files:
#             results["message"] = "No files found in incoming/"
#             return results

#         librarian = LibrarianAgent()

#         for file_path in incoming_files:
#             filename = file_path.split("/")[-1]
#             if not filename:
#                 continue  # skip folder placeholders

#             # ------ TEXT EXTRACTION ------
#             text = extract_text_from_blob(bucket, file_path)

#             if not text.strip():
#                 # empty or unreadable → move to Others and mark skipped
#                 results["skipped"] += 1
#                 dest = f"Others/{filename}"
#                 provider_obj.move_file(file_path, dest)
#                 results["details"].append(
#                     {
#                         "file": file_path,
#                         "type": "EMPTY",
#                         "moved_to": dest,
#                         "status": "MOVED",
#                     }
#                 )
#                 continue

#             # ------ CLASSIFICATION (Gemini) ------
#             # LibrarianAgent.classify_document() returns a *string*:
#             # "RESUME", "JD", "REFERENCE", "UNKNOWN"
#             doc_type = librarian.classify_document(text)

#             dest = ""
#             if doc_type == "RESUME":
#                 dest = f"Resumes/{filename}"
#                 cand_id = f"CAND-{uuid.uuid4().hex[:8]}"
#                 save_candidate(cand_id, filename, text)
#                 results["resumes"] += 1

#             elif doc_type == "JD":
#                 dest = f"JDs/{filename}"
#                 jd_id = f"JD-{uuid.uuid4().hex[:8]}"
#                 save_jd(jd_id, filename, text)
#                 results["jds"] += 1

#             elif doc_type in ("REFERENCE", "UNKNOWN"):
#                 dest = f"Others/{filename}"
#                 results["unknown"] += 1

#             else:
#                 # fallback safety
#                 dest = f"Others/{filename}"
#                 results["others"] += 1

#             # ------ MOVE FILE ------
#             provider_obj.move_file(file_path, dest)

#             results["details"].append(
#                 {
#                     "file": file_path,
#                     "type": doc_type,
#                     "moved_to": dest,
#                     "status": "MOVED",
#                 }
#             )

#         return results

#     except Exception as e:
#         print("\n[PROCESS FILES ERROR]")
#         traceback.print_exc()
#         raise HTTPException(
#             status_code=500,
#             detail=f"Error processing files: {str(e)}",
#         )


# # ---------------------------------------------------------
# # DATA RETRIEVAL ROUTES (for Workspace)
# # ---------------------------------------------------------
# @app.get("/jds")
# def get_jds_route():
#     try:
#         return {"jds": list_jds()}
#     except Exception as e:
#         print("[ERROR] listing JDs:", e)
#         return {"jds": []}


# @app.get("/candidates")
# def get_candidates_route():
#     try:
#         return {"candidates": list_candidates()}
#     except Exception as e:
#         print("[ERROR] listing candidates:", e)
#         return {"candidates": []}


# # ---------------------------------------------------------
# # MAPPING ROUTES
# # ---------------------------------------------------------
# @app.post("/map-candidate-to-jd")
# def map_candidate_to_jd(payload: MappingRequest):
#     if not get_jd(payload.jd_id):
#         raise HTTPException(404, "JD not found")
#     if not get_candidate(payload.candidate_id):
#         raise HTTPException(404, "Candidate not found")

#     save_mapping(payload.jd_id, payload.candidate_id, payload.status)
#     return {"message": "Mapping saved"}


# @app.get("/jd/{jd_id}/candidates")
# def get_candidates_for_jd(jd_id: str):
#     """
#     Returns all candidates mapped to this JD.
#     Used by the workspace right panel.
#     """
#     mappings = list_mappings_for_jd(jd_id)
#     candidate_ids = [m["candidate_id"] for m in mappings]
#     all_candidates = {c["candidate_id"]: c for c in list_candidates()}
#     linked = [all_candidates[cid] for cid in candidate_ids if cid in all_candidates]

#     return {
#         "jd_id": jd_id,
#         "candidates": linked,
#         "mappings": mappings,
#     }


# # ---------------------------------------------------------
# # EVALUATION / SHORTLIST ROUTES
# # ---------------------------------------------------------
# @app.post("/candidate/evaluate")
# def evaluate_candidate_api(payload: CandidateEvalRequest):
#     """
#     Generic evaluation route (not used by current UI but kept for testing).
#     """
#     return evaluate_candidate(
#         job_id=payload.job_id,
#         candidate_id=payload.candidate_id,
#         required_experience=payload.required_experience,
#         candidate_experience=payload.candidate_experience,
#         career_gaps_detected=payload.career_gaps_detected,
#         job_shifts_last_year=payload.job_shifts_last_year,
#         skills_match_score=payload.skills_match_score,
#     )


# @app.post("/jd/{jd_id}/shortlist")
# def shortlist_candidate(jd_id: str, payload: ShortlistRequest):
#     """
#     Used by the frontend 'Evaluate Agent' button.
#     For now we use fixed experience values and rules from recruiter.py.
#     """
#     jd = get_jd(jd_id)
#     cand = get_candidate(payload.candidate_id)

#     if not jd or not cand:
#         raise HTTPException(404, "JD or Candidate not found")

#     eval_result = evaluate_candidate(
#         job_id=jd_id,
#         candidate_id=payload.candidate_id,
#         required_experience=3.0,      # TODO: derive from JD text
#         candidate_experience=3.0,     # TODO: derive from resume text
#         career_gaps_detected=False,   # TODO: derive from resume timeline
#         job_shifts_last_year=0,       # TODO: derive from resume history
#         skills_match_score=payload.skills_match_score,
#     )

#     save_shortlist_result(jd_id, payload.candidate_id, eval_result)
#     return eval_result

from dotenv import load_dotenv
import os
import uuid
import traceback
from io import BytesIO
from typing import Literal
from datetime import timedelta

from fastapi import FastAPI, Query, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi import Query

# ---------- EXTERNAL IMPORTS ----------
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

# ---------- INTERNAL IMPORTS ----------
from backend.db.cloud_provider import (
    AWSProvider,
    GCPProvider,
    AzureProvider,
    CloudProvider,
)
from backend.agents.recruiter import evaluate_candidate
from backend.agents.cloud_setup import CloudSetupAgent
from backend.agents.librarian import LibrarianAgent

# ---------- NEW IMPORTS FOR AZURE FLOW ----------
import os
import httpx
from pydantic import BaseModel

from backend.cloud_config import (
    save_azure_config,
    get_azure_config,
    delete_azure_config,
    set_azure_status,
    get_azure_status,
)

from backend.db.firestore_client import (
    save_jd,
    save_candidate,
    list_jds,
    list_candidates,
    save_mapping,
    list_mappings_for_jd,
    get_candidate,
    get_jd,
    save_shortlist_result,
)

load_dotenv()

# ---------------------------------------------------------
# FASTAPI APP
# ---------------------------------------------------------
app = FastAPI(title="CARPAS – Cloud-Agnostic Recruitment System")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------
# INTERNAL AZURE CONFIG UPDATE (called by GitHub Action)
# ---------------------------------------------------------
INTERNAL_WEBHOOK_SECRET = os.getenv("INTERNAL_WEBHOOK_SECRET")


@app.post("/internal/cloud/azure/config")
async def internal_update_azure_config(
    payload: AzureConfigPayload,
    x_internal_secret: str = Header(default=None, alias="X-Internal-Secret"),
):
    """
    PRIVATE: Called only from GitHub Action after Terraform finishes.
    Saves the Azure connection string & container name into Firestore.
    """

    if not INTERNAL_WEBHOOK_SECRET or x_internal_secret != INTERNAL_WEBHOOK_SECRET:
        # Protect this endpoint from external calls
        raise HTTPException(status_code=403, detail="Unauthorized")

    save_azure_config(
        env=payload.environment,
        connection_string=payload.connection_string,
        container_name=payload.container_name,
    )

    return {"status": "ok"}


@app.post("/internal/cloud/azure/config/destroyed")
async def internal_azure_destroyed(
    payload: AzureDestroyPayload,
    x_internal_secret: str = Header(default=None, alias="X-Internal-Secret")
):
    if not INTERNAL_WEBHOOK_SECRET or x_internal_secret != INTERNAL_WEBHOOK_SECRET:
        raise HTTPException(status_code=403, detail="Unauthorized")

    delete_azure_config(payload.environment)
    return {"status": "deleted"}


# ---------------------------------------------------------
# PUBLIC AZURE PROVISIONING ENDPOINT (triggered from frontend)
# ---------------------------------------------------------

GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_REPO = os.getenv("GITHUB_REPO")  # e.g. "your-org/TAG"


@app.post("/cloud/azure/provision")
async def provision_azure(req: UseAzureRequest):
    """
    Trigger Azure infrastructure provisioning via GitHub Actions.
    - If Azure config already exists for this environment -> return 'ready'
    - Else -> trigger repo_dispatch 'provision-azure' and return 'provisioning'
    """

    # 1. Check if already provisioned
    existing = get_azure_config(req.environment)
    if existing and existing.get("status") == "ready":
        return {
            "status": "ready",
            "message": f"Azure already provisioned for env={req.environment}",
        }

    # Mark as provisioning in Firestore
    set_azure_status(req.environment, "provisioning")

    # 2. Validate GitHub config
    if not GITHUB_TOKEN or not GITHUB_REPO:
        raise HTTPException(
            status_code=500,
            detail="GitHub integration not configured (GITHUB_TOKEN / GITHUB_REPO missing)",
        )

    # 3. Trigger GitHub repo dispatch to start the Terraform workflow
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"https://api.github.com/repos/{GITHUB_REPO}/dispatches",
            headers={
                "Accept": "application/vnd.github+json",
                "Authorization": f"Bearer {GITHUB_TOKEN}",
            },
            json={
                "event_type": "provision-azure",
                "client_payload": {
                    "environment": req.environment,
                },
            },
        )

    if resp.status_code not in (200, 201, 204):
        print("❌ GitHub dispatch error:", resp.status_code, resp.text)
        raise HTTPException(
            status_code=500,
            detail="Failed to trigger Azure provisioning via GitHub Actions",
        )


    return {
        "status": "provisioning",
        "message": f"Azure provisioning started for env={req.environment}",
    }


@app.post("/cloud/azure/destroy")
async def destroy_azure(req: UseAzureRequest):
    """
    Trigger Azure destroy workflow.
    """
    # If no config → nothing to destroy
    existing = get_azure_config(req.environment)
    if not existing:
        return {
            "status": "not_provisioned",
            "message": f"No Azure infra found for env={req.environment}",
        }

    # Mark status as deleting
    set_azure_status(req.environment, "deleting")

    # Trigger destroy workflow
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            f"https://api.github.com/repos/{GITHUB_REPO}/dispatches",
            headers={
                "Accept": "application/vnd.github+json",
                "Authorization": f"Bearer {GITHUB_TOKEN}",
            },
            json={
                "event_type": "destroy-azure",
                "client_payload": {
                    "environment": req.environment,
                },
            },
        )

    if resp.status_code not in (200, 201, 204):
        raise HTTPException(
            status_code=500,
            detail="Failed to trigger Azure destroy workflow"
        )

    return {
        "status": "destroying",
        "message": f"Azure destruction started for env={req.environment}"
    }





# ---------------------------------------------------------
# MODELS
# ---------------------------------------------------------
CloudName = Literal["aws", "gcp", "azure"]

# ---------------------------------------------------------
# AZURE CLOUD CONFIG MODELS
# ---------------------------------------------------------
class AzureConfigPayload(BaseModel):
    environment: str
    connection_string: str
    container_name: str

class AzureDestroyPayload(BaseModel):
    environment: str


class UseAzureRequest(BaseModel):
    # You can later switch this to a project-specific id if needed
    environment: str = "dev"

class UploadURLRequest(BaseModel):
    filename: str

class CandidateEvalRequest(BaseModel):
    job_id: str
    candidate_id: str
    required_experience: float
    candidate_experience: float
    career_gaps_detected: bool
    job_shifts_last_year: int
    skills_match_score: int

class MappingRequest(BaseModel):
    jd_id: str
    candidate_id: str
    status: Literal["PENDING", "SHORTLISTED", "REJECTED"] = "PENDING"

class ShortlistRequest(BaseModel):
    candidate_id: str
    skills_match_score: int = 70

# ---------------------------------------------------------
# HELPERS
# ---------------------------------------------------------
def get_provider(provider: CloudName) -> CloudProvider:
    if provider == "gcp":
        return GCPProvider()
    if provider == "aws":
        return AWSProvider()
    if provider == "azure":
        return AzureProvider()
    raise HTTPException(400, "Invalid provider")

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    try:
        ext = filename.lower().split(".")[-1]

        if ext == "pdf":
            return pdf_extract_text(BytesIO(file_bytes))

        if ext == "docx":
            try:
                doc = Document(BytesIO(file_bytes))
                return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            except:
                return ""

        if ext in ("txt", "md"):
            return file_bytes.decode("utf-8", errors="ignore")

        return ""
    except Exception as e:
        print("TEXT EXTRACTION ERROR:", e)
        return ""

# ---------------------------------------------------------
# CLOUD SETUP
# ---------------------------------------------------------
@app.get("/cloud/setup")
def setup_cloud(provider: CloudName):
    provider_obj = get_provider(provider)
    agent = CloudSetupAgent(provider_obj)

    session_id = uuid.uuid4().hex
    response = agent.setup_storage()
    response["session_id"] = session_id

    return response

# ---------------------------------------------------------
# SIGNED UPLOAD URL
# ---------------------------------------------------------
@app.post("/cloud/generate-upload-url")
def generate_upload_url(payload: UploadURLRequest, provider: CloudName = Query(...)):
    try:
        provider_obj = get_provider(provider)
        object_path = f"incoming/{payload.filename}"

        # FIXED: Correct signature method
        url = provider_obj.generate_upload_link(object_path)

        return {"upload_url": url, "object_path": object_path}

    except Exception as e:
        print("UPLOAD URL ERROR:", e)
        raise HTTPException(500, str(e))

# ---------------------------------------------------------
# PROCESS FILES
# ---------------------------------------------------------
@app.post("/cloud/process-files")
def process_files(provider: CloudName = Query(...), session_id: str = Query(...)):
    from backend.agents.jd_extractor import extract_jd_title
    from backend.agents.name_extractor import extract_name_from_text
    from backend.agents.skills_extractor import extract_skills

    try:
        provider_obj = get_provider(provider)
        incoming_files = provider_obj.list_files("incoming/")

        librarian = LibrarianAgent()
        results = []

        for file_path in incoming_files:
            filename = os.path.basename(file_path)

            # 1. Download bytes (provider agnostic)
            try:
                file_bytes = provider_obj.download_bytes(file_path)
            except:
                provider_obj.move_file(file_path, f"Others/{filename}")
                continue

            # 2. Extract text
            text = extract_text_from_bytes(file_bytes, filename)
            if not text:
                provider_obj.move_file(file_path, f"Others/{filename}")
                continue

            # 3. Classify
            doc_type = librarian.classify_document(text)

            # ---------------- JD ----------------
            if doc_type == "JD":
                jd_id = f"JD-{uuid.uuid4().hex[:8]}"
                job_title = extract_jd_title(text)

                save_jd(jd_id, filename, text, session_id, job_title)
                provider_obj.move_file(file_path, f"JDs/{filename}")

                results.append({
                    "file": filename,
                    "type": "JD",
                    "jd_id": jd_id,
                    "job_title": job_title
                })
                continue

            # ---------------- RESUME ----------------
            if doc_type == "RESUME":
                cand_id = f"CAND-{uuid.uuid4().hex[:8]}"
                name = extract_name_from_text(text)
                skills = extract_skills(text)

                save_candidate(cand_id, filename, text, session_id, name, skills)
                provider_obj.move_file(file_path, f"Resumes/{filename}")

                results.append({
                    "file": filename,
                    "type": "RESUME",
                    "candidate_id": cand_id,
                    "candidate_name": name,
                    "skills": skills
                })
                continue

            # ---------------- OTHER ----------------
            provider_obj.move_file(file_path, f"Others/{filename}")
            results.append({"file": filename, "type": "OTHER"})

        return {
            "message": "Processing completed",
            "session_id": session_id,
            "details": results
        }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(500, str(e))

# ---------------------------------------------------------
# DATA FETCH
# ---------------------------------------------------------
@app.get("/jds")
def get_jds_route(session_id: str):
    return {"jds": list_jds(session_id)}

@app.get("/candidates")
def get_candidates_route(session_id: str):
    return {"candidates": list_candidates(session_id)}

# ---------------------------------------------------------
# MAPPING
# ---------------------------------------------------------
@app.post("/map-candidate-to-jd")
def map_candidate_to_jd(payload: MappingRequest):
    if not get_jd(payload.jd_id):
        raise HTTPException(404, "JD not found")
    if not get_candidate(payload.candidate_id):
        raise HTTPException(404, "Candidate not found")

    save_mapping(payload.jd_id, payload.candidate_id, payload.status)
    return {"message": "Mapping saved"}

@app.get("/jd/{jd_id}/candidates")
def get_candidates_for_jd(jd_id: str, session_id: str):
    mapped = list_mappings_for_jd(jd_id)
    mapped_ids = {m["candidate_id"] for m in mapped}

    all_candidates = {c["candidate_id"]: c for c in list_candidates(session_id)}

    mapped_candidates = [
        all_candidates[cid] for cid in mapped_ids if cid in all_candidates
    ]

    return {"candidates": mapped_candidates, "mappings": mapped}

# ---------------------------------------------------------
# EVALUATION
# ---------------------------------------------------------
@app.post("/candidate/evaluate")
def evaluate_candidate_api(payload: CandidateEvalRequest):
    return evaluate_candidate(**payload.dict())

@app.post("/jd/{jd_id}/shortlist")
def shortlist_candidate(jd_id: str, payload: ShortlistRequest):
    jd = get_jd(jd_id)
    cand = get_candidate(payload.candidate_id)

    if not jd or not cand:
        raise HTTPException(404, "Not found")

    result = evaluate_candidate(
        job_id=jd_id,
        candidate_id=payload.candidate_id,
        required_experience=3.0,
        candidate_experience=3.0,
        career_gaps_detected=False,
        job_shifts_last_year=0,
        skills_match_score=payload.skills_match_score,
    )

    save_shortlist_result(jd_id, payload.candidate_id, result)
    return result

@app.get("/cloud/azure/status")
async def get_azure_status_api(environment: str = Query("dev")):
    """
    Return the current Azure status for this environment.
    Values can be:
      - not_provisioned
      - provisioning
      - ready
      - deleting
      - unknown
    """
    status = get_azure_status(environment)
    return {"environment": environment, "status": status}